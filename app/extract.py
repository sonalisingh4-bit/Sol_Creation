"""Extract plain text from uploaded source files (pdf / docx / txt / images)."""
from __future__ import annotations

from pathlib import Path

from pypdf import PdfReader

from . import gemini_client

TEXT_EXTS = {".txt", ".md", ".csv"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff"}
DOC_EXTS = {".docx"}
PDF_EXTS = {".pdf"}
SUPPORTED_EXTS = TEXT_EXTS | IMAGE_EXTS | DOC_EXTS | PDF_EXTS

# Below this many chars per page we assume the PDF is scanned and OCR via Gemini.
_MIN_CHARS_PER_PAGE = 50

_OCR_PROMPT = (
    "Transcribe ALL text content from this document verbatim. Preserve headings, "
    "lists, tables (as readable text), and equations. Do not summarise or add commentary."
)


def is_supported(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_EXTS


def _extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            pages.append("")
    return "\n\n".join(pages).strip()


def _extract_docx(path: Path) -> str:
    from docx import Document  # imported lazily

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _ocr_generate(uploaded) -> str:
    """Use Gemini multimodal to read a scanned PDF or an image."""
    from .gemini_client import get_client
    from . import config
    from google.genai import types

    client = get_client()
    resp = client.models.generate_content(
        model=config.GEMINI_PARSE_MODEL,
        contents=[uploaded, _OCR_PROMPT],
        config=types.GenerateContentConfig(temperature=0.0),
    )
    return (resp.text or "").strip()


def extract_text(path: str | Path) -> str:
    """Return extracted plain text for a supported file, OCR'ing scans/images."""
    path = Path(path)
    ext = path.suffix.lower()

    if ext in TEXT_EXTS:
        return path.read_text(encoding="utf-8", errors="ignore").strip()

    if ext in DOC_EXTS:
        return _extract_docx(path)

    if ext in IMAGE_EXTS:
        return _ocr_generate(gemini_client.upload_file(path))

    if ext in PDF_EXTS:
        text = _extract_pdf_text(path)
        n_pages = max(len(PdfReader(str(path)).pages), 1)
        if len(text) < _MIN_CHARS_PER_PAGE * n_pages:
            # Looks scanned -> OCR the whole file via Gemini.
            ocr = _ocr_generate(gemini_client.upload_file(path))
            if len(ocr) > len(text):
                return ocr
        return text

    raise ValueError(f"Unsupported file type: {ext}")
