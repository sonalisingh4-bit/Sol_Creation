"""Render a SolvedPaper into a .docx (and, if possible, a .pdf)."""
from __future__ import annotations

import html
import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from . import figures, mathrender, pdf_convert
from .solver import SolvedItem, SolvedPaper, SolvedQuestion

# Nirmala UI ships with Windows and covers Latin + all major Indic scripts,
# so a single font renders every supported language correctly.
_FONT = "Nirmala UI"
_GREY = RGBColor(0x66, 0x66, 0x66)
_ACCENT = RGBColor(0x1A, 0x47, 0x8A)

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_BULLET_RE = re.compile(r"^\s*[-*•]\s+(.*)")
_NUM_RE = re.compile(r"^\s*(\d+)[.)]\s+(.*)")
_HEAD_RE = re.compile(r"^\s*#{1,6}\s+(.*)")

# Defensive cleanup: models sometimes emit HTML (huge &nbsp; runs, <br>, even a
# hallucinated <img src=...>) when asked to "draw" a structure. None of that
# belongs in the document, so strip it before rendering.
_IMG_RE = re.compile(r"<img[^>]*>", re.IGNORECASE)
_BR_RE = re.compile(r"<br\s*/?>", re.IGNORECASE)
_TAG_RE = re.compile(r"</?[a-zA-Z][^>]*>")

# Models sometimes emit LaTeX math (\xrightarrow, \frac, $...$) in an otherwise
# plain-text answer. Convert the common bits to readable symbols instead of showing
# raw markup like "2KMnO4 $\xrightarrow{513 K}$ ...".
_LTX_ARROW = re.compile(r"\\xrightarrow\s*(?:\[[^\]]*\])?\s*\{([^}]*)\}")
_LTX_FRAC = re.compile(r"\\frac\s*\{([^}]*)\}\s*\{([^}]*)\}")
_LTX_TEXT = re.compile(r"\\(?:text|mathrm|mathbf|mathit|mathsf)\s*\{([^}]*)\}")
_LTX_INLINE = re.compile(r"\$([^$]*)\$")
_LTX_WORDS = {
    r"\Delta": "Δ", r"\alpha": "α", r"\beta": "β", r"\gamma": "γ", r"\pi": "π",
    r"\rightleftharpoons": " ⇌ ", r"\leftrightarrow": " ⇌ ",
    r"\rightarrow": "→", r"\to": "→", r"\times": "×", r"\cdot": "·", r"\pm": "±",
    r"\Rightarrow": "⇒", r"\left": "", r"\right": "", r"\,": " ", r"\ ": " ",
}


def _read_braced(text: str, start: int) -> tuple[str, int] | None:
    if start >= len(text) or text[start] != "{":
        return None
    depth = 0
    for i in range(start, len(text)):
        ch = text[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return text[start + 1 : i], i + 1
    return None


def _skip_ws(text: str, i: int) -> int:
    while i < len(text) and text[i].isspace():
        i += 1
    return i


def _replace_latex_groups(text: str) -> str:
    r"""Readable fallback for nested LaTeX like \frac{\sqrt{x}}{7}."""
    out: list[str] = []
    i = 0
    while i < len(text):
        if text.startswith(r"\frac", i):
            j = _skip_ws(text, i + len(r"\frac"))
            num = _read_braced(text, j)
            if num is not None:
                den = _read_braced(text, _skip_ws(text, num[1]))
                if den is not None:
                    out.append(
                        f"({_replace_latex_groups(num[0])})/({_replace_latex_groups(den[0])})"
                    )
                    i = den[1]
                    continue
        if text.startswith(r"\sqrt", i):
            j = _skip_ws(text, i + len(r"\sqrt"))
            if j < len(text) and text[j] == "[":
                end = text.find("]", j + 1)
                if end != -1:
                    j = _skip_ws(text, end + 1)
            body = _read_braced(text, j)
            if body is not None:
                out.append(f"√({_replace_latex_groups(body[0])})")
                i = body[1]
                continue
        out.append(text[i])
        i += 1
    return "".join(out)


def _clean_ad_hoc_math(text: str) -> str:
    """Repair common non-LaTeX math fragments so raw tokens never reach students."""
    text = re.sub(
        r"\bint_([^\s^]+)\^\(([^)]+)\)",
        lambda m: f"∫_{m.group(1)}^({m.group(2)})",
        text,
    )
    text = re.sub(
        r"\bint_([^\s^]+)\^\{([^}]+)\}",
        lambda m: f"∫_{m.group(1)}^({m.group(2)})",
        text,
    )
    text = re.sub(r"\bint_", "∫_", text)
    for fn in ("sin", "cos", "tan", "cot", "sec", "cosec"):
        text = re.sub(rf"\b{fn}\^\(-1\)", f"{fn}⁻¹", text)
        text = re.sub(rf"\b{fn}\^-1", f"{fn}⁻¹", text)
    text = re.sub(
        r"\bfrac\s*√\s*([A-Za-z0-9]+?)([0-9])\b",
        lambda m: f"(√{m.group(1)})/({m.group(2)})",
        text,
    )
    text = re.sub(
        r"\bfrac\s*([A-Za-z0-9]+)\s*/\s*([A-Za-z0-9]+)\b",
        lambda m: f"({m.group(1)})/({m.group(2)})",
        text,
    )
    text = re.sub(
        r"(sin⁻¹|cos⁻¹|tan⁻¹)\s*√\(([^()]+)\)/\(([^()]+)\)",
        lambda m: f"{m.group(1)}√({m.group(2)}/({m.group(3)}))",
        text,
    )
    return text


def _delatex(text: str) -> str:
    text = _replace_latex_groups(text)
    text = _LTX_ARROW.sub(lambda m: f" →({m.group(1).strip()}) " if m.group(1).strip() else " → ", text)
    text = _LTX_FRAC.sub(r"(\1)/(\2)", text)
    text = _LTX_TEXT.sub(r"\1", text)
    for k, v in _LTX_WORDS.items():
        text = text.replace(k, v)
    for fn in ("sin", "cos", "tan", "cot", "sec", "cosec", "log", "ln"):
        text = text.replace(rf"\{fn}^{{-1}}", f"{fn}⁻¹")
        text = text.replace(rf"\{fn}", fn)
    text = _LTX_INLINE.sub(r"\1", text)  # drop $...$ delimiters, keep the content
    return _clean_ad_hoc_math(text)


# --- Math ------------------------------------------------------------------
# The model writes mathematics as LaTeX: $...$ / \(...\) inline, $$...$$ / \[...\]
# on its own line. We render each fragment as a real Word equation (OMML). Split
# a line into alternating prose / math so prose stays text and math becomes an
# equation. Longest delimiters first so "$$" wins over "$".
_MATH_SEG = re.compile(
    r"\$\$(?P<d1>.+?)\$\$"        # $$ display $$
    r"|\\\[(?P<d2>.+?)\\\]"       # \[ display \]
    r"|\$(?P<i1>.+?)\$"           # $ inline $
    r"|\\\((?P<i2>.+?)\\\)",      # \( inline \)
    re.DOTALL,
)

# Extra symbol map used only for the text FALLBACK (machine without Word/Office):
# turn the commonest LaTeX into readable Unicode so even the fallback is legible.
_LTX_FALLBACK = {
    r"\sqrt": "√", r"\cdot": "·", r"\times": "×", r"\pm": "±", r"\mp": "∓",
    r"\leq": "≤", r"\geq": "≥", r"\neq": "≠", r"\approx": "≈", r"\equiv": "≡",
    r"\Rightarrow": "⇒", r"\Leftarrow": "⇐", r"\rightarrow": "→", r"\to": "→",
    r"\infty": "∞", r"\theta": "θ", r"\alpha": "α", r"\beta": "β", r"\gamma": "γ",
    r"\lambda": "λ", r"\mu": "μ", r"\pi": "π", r"\phi": "φ", r"\omega": "ω",
    r"\Delta": "Δ", r"\circ": "°", r"\left": "", r"\right": "", r"\,": " ", r"\;": " ",
}
_LTX_CMD = re.compile(r"\\(?:frac|sqrt|vec|hat|bar|overline|mathrm|mathbf|text|left|right)\b")
_LTX_ANYCMD = re.compile(r"\\([a-zA-Z]+)")


def _latex_to_text(latex: str) -> str:
    """Readable plain-text rendering of a LaTeX fragment — the fallback used only
    when native equation rendering is unavailable or a fragment won't convert."""
    s = _replace_latex_groups(latex)
    s = _LTX_FRAC.sub(r"(\1)/(\2)", s)
    for k, v in _LTX_FALLBACK.items():
        s = s.replace(k, v)
    s = re.sub(r"\^\{([^{}]*)\}", r"^(\1)", s)   # ^{...} -> ^(...)
    s = re.sub(r"_\{([^{}]*)\}", r"_\1", s)       # _{...} -> _...
    s = re.sub(r"\\[a-zA-Z]+\{([^{}]*)\}", r"\1", s)  # \hat{i} etc -> i
    s = _LTX_ANYCMD.sub(r"\1", s)                  # \sin -> sin, drop stray commands
    return _clean_ad_hoc_math(s.replace("{", "").replace("}", "")).strip()


def _join_multiline_math(text: str) -> str:
    r"""Put each $$...$$ / \[...\] / \(...\) block on a single line.

    The document is rendered line by line, so a display equation the model split
    across several lines would never be recognised as one block — its content would
    leak out as prose and the equation would look dropped. Collapse the internal
    newlines first so the block stays intact."""
    def _collapse(m: re.Match) -> str:
        return re.sub(r"\s*\n\s*", " ", m.group(0))

    text = re.sub(r"\$\$.+?\$\$", _collapse, text, flags=re.DOTALL)
    text = re.sub(r"\\\[.+?\\\]", _collapse, text, flags=re.DOTALL)
    text = re.sub(r"\\\(.+?\\\)", _collapse, text, flags=re.DOTALL)
    return text


def _iter_math_segments(text: str):
    """Yield ('text', str) / ('math', latex) segments of a line, in order."""
    pos = 0
    for m in _MATH_SEG.finditer(text):
        if m.start() > pos:
            yield "text", text[pos : m.start()]
        yield "math", (m.group("d1") or m.group("d2") or m.group("i1") or m.group("i2"))
        pos = m.end()
    if pos < len(text):
        yield "text", text[pos:]


def _append_equation(paragraph, latex: str) -> bool:
    """Append a native Word equation to `paragraph`. Returns False if it could not
    be rendered (no Office stylesheet, or LaTeX the converter rejects)."""
    el = mathrender.latex_to_omath(latex)
    if el is None:
        png = mathrender.latex_to_png(latex)
        if png is None:
            return False
        # Preserve the PNG's DPI-based natural size. Forcing every formula to the
        # same height makes simple symbols look huge and tall fractions look tiny.
        paragraph.add_run().add_picture(BytesIO(png))
        return True
    paragraph._p.append(el)
    return True


def _clean_answer(answer: str) -> str:
    if not answer:
        return ""
    text = _IMG_RE.sub("", answer)        # drop hallucinated image embeds
    text = _BR_RE.sub("\n", text)         # <br> -> real line break
    text = _TAG_RE.sub("", text)          # strip any other HTML tags
    text = html.unescape(text)            # &nbsp; -> \xa0, &amp; -> & ...
    # NOTE: math ($...$) is deliberately preserved here and rendered as real Word
    # equations later (see _add_rich_line). Per-segment cleanup handles stray LaTeX.
    text = text.replace("\xa0", " ")      # non-breaking space -> normal space
    text = re.sub(r"[ \t]{2,}", " ", text)  # collapse the &nbsp; "spacing" floods
    text = re.sub(r"\n{3,}", "\n\n", text)  # collapse big vertical gaps
    text = _join_multiline_math(text)     # keep each display equation on one line
    return "\n".join(ln.rstrip() for ln in text.splitlines())


def _style_run(run, *, size=11, bold=False, italic=False, color=None) -> None:
    run.font.name = _FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), _FONT)


def _add_text_runs(paragraph, text: str, *, size=11, color=None) -> None:
    """Add prose runs, honouring **bold** spans."""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            _style_run(paragraph.add_run(text[pos : m.start()]), size=size, color=color)
        _style_run(paragraph.add_run(m.group(1)), size=size, bold=True, color=color)
        pos = m.end()
    if pos < len(text):
        _style_run(paragraph.add_run(text[pos:]), size=size, color=color)


def _add_rich_line(paragraph, text: str, *, size=11, color=None) -> None:
    """Render a single line: prose as styled runs (with **bold**), and each LaTeX
    fragment ($...$, $$...$$) as a native Word equation — inline with the prose."""
    for kind, content in _iter_math_segments(text):
        if kind == "math":
            if _append_equation(paragraph, content):
                continue
            # No native rendering available — degrade to readable text, not raw LaTeX.
            _add_text_runs(paragraph, _latex_to_text(content), size=size, color=color)
        else:
            # Clean any stray LaTeX the model left outside $...$ (arrows, \frac …). A
            # lone '$' here is an unpaired delimiter, never prose — drop it so it can
            # never surface as a literal dollar sign in the document.
            _add_text_runs(paragraph, _delatex(content).replace("$", ""), size=size, color=color)


def _embed_png(doc: Document, png: bytes, width, caption: str = "") -> None:
    # Keep a figure with the label that introduces it ("G:", "A:" …) so a page
    # break never strands the label at the foot of one page with its picture at
    # the top of the next — which reads to the user as "the diagram is missing".
    paras = doc.paragraphs
    if paras:
        paras[-1].paragraph_format.keep_with_next = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p.paragraph_format.keep_with_next = True  # keep the picture with its caption
    p.add_run().add_picture(BytesIO(png), width=width)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_run(cap.add_run(caption), size=9, italic=True, color=_GREY)


# --- Markdown tables ------------------------------------------------------
# The model emits GitHub-flavoured tables ("| a | b |" + "| :--- | :--- |") for
# comparisons (bleaching action, physisorption vs chemisorption, extraction …).
# Rendered line-by-line these show ugly raw "| :--- |" dashes, so detect the
# block and build a real Word table instead. Cells may span several physical
# lines (equations stacked in one cell), so a row is only "complete" once it has
# collected ncols+1 pipes and ends with a pipe.
def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.count("|") >= 2


def _is_table_sep(line: str) -> bool:
    s = line.strip()
    return "-" in s and set(s) <= set("|:- ")


def _split_row(row: str) -> list[str]:
    parts = row.strip().split("|")
    if parts and parts[0].strip() == "":
        parts = parts[1:]
    if parts and parts[-1].strip() == "":
        parts = parts[:-1]
    return [p.strip() for p in parts]


def _emit_table(doc: Document, rows: list[list[str]], ncols: int) -> None:
    table = doc.add_table(rows=len(rows), cols=ncols)
    try:
        table.style = "Table Grid"  # built-in style that draws cell borders
    except Exception:  # noqa: BLE001 - unusual template without the style
        pass
    for r, cells in enumerate(rows):
        for c in range(ncols):
            cell = table.cell(r, c)
            txt = cells[c] if c < len(cells) else ""
            first = cell.paragraphs[0]
            for k, sub in enumerate(txt.split("\n")):
                p = first if k == 0 else cell.add_paragraph()
                _add_rich_line(p, sub.strip(), size=10)
                if r == 0:  # header row in bold
                    for run in p.runs:
                        run.font.bold = True


def _render_markdown_table(doc: Document, lines: list[str], start: int) -> int:
    """Render the table beginning at lines[start]; return the next line index."""
    ncols = len(_split_row(lines[start]))
    rows: list[list[str]] = [_split_row(lines[start])]
    i, n = start + 2, len(lines)  # skip the header and the "| :--- |" separator
    buf = ""
    while i < n:
        ln = lines[i].rstrip()
        if not buf:
            if not ln.strip().startswith("|"):
                break  # first line that is not part of the table ends it
            buf = ln.strip()
        else:
            buf += "\n" + ln  # continuation of a multi-line cell
        if buf.count("|") >= ncols + 1 and buf.rstrip().endswith("|"):
            rows.append(_split_row(buf))
            buf = ""
        i += 1
    if buf.strip():  # a trailing row that never closed its final pipe
        rows.append(_split_row(buf))
    if ncols >= 1:
        _emit_table(doc, rows, ncols)
    return i


def _render_text_block(doc: Document, text: str) -> None:
    """Light markdown rendering: tables, headings, bullet/numbered lists, bold, paragraphs."""
    lines = text.splitlines()
    i, n = 0, len(lines)
    while i < n:
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        # Markdown table: a "| … |" row immediately followed by a "| :--- | … |"
        # separator. Render it as a real Word table, not raw pipes/dashes.
        if _is_table_row(line) and i + 1 < n and _is_table_sep(lines[i + 1]):
            i = _render_markdown_table(doc, lines, i)
            continue

        # Safety net: a bare SMILES the model forgot to wrap in [[FIG]] still draws.
        # Skip lines carrying math ($...$) — LaTeX must never be parsed as SMILES.
        if "$" not in line:
            png = figures.smiles_line_png(line)
            if png is not None:
                _embed_png(doc, png, Inches(figures.WIDTH["MOL"]))
                i += 1
                continue

        if m := _HEAD_RE.match(line):
            p = doc.add_paragraph()
            _add_rich_line(p, m.group(1), size=12)
            for r in p.runs:
                r.font.bold = True
            i += 1
            continue
        if m := _BULLET_RE.match(line):
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_line(p, m.group(1))
            i += 1
            continue
        if m := _NUM_RE.match(line):
            # Keep the model's own number (it restarts at 1 within each answer) as
            # literal text. Word's "List Number" style shares ONE counter across the
            # whole document, so separate lists would otherwise run on (1, 2 … 36)
            # instead of restarting per question.
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Inches(0.4)
            pf.first_line_indent = Inches(-0.25)  # hanging indent, like a real list
            _add_rich_line(p, f"{m.group(1)}. {m.group(2)}")
            i += 1
            continue

        p = doc.add_paragraph()
        _add_rich_line(p, line)
        i += 1


def _embed_directive(doc: Document, match: re.Match) -> None:
    """Render a [[FIG ...]] directive as a real image, or fall back to caption text."""
    png, kind, caption, fallback = figures.render_match(match)
    if png is None:
        # Rendering unavailable or spec invalid — keep a text fallback so the answer
        # never silently loses information.
        if fallback:
            p = doc.add_paragraph()
            _add_rich_line(p, fallback)
        return
    _embed_png(doc, png, Inches(figures.WIDTH.get(kind, 4.5)), caption)


def _add_answer(doc: Document, answer: str) -> None:
    """Render the answer, embedding figure directives as real images."""
    text = _clean_answer(answer)
    pos = 0
    for m in figures.DIRECTIVE_RE.finditer(text):
        _render_text_block(doc, text[pos : m.start()])
        _embed_directive(doc, m)
        pos = m.end()
    _render_text_block(doc, text[pos:])


def _heading(doc: Document, number: str, marks, text: str) -> None:
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    label = f"Q{number}" if not number.lower().startswith("q") else number
    if marks is not None:
        label += f"   [{marks} marks]" if marks != 1 else f"   [{marks} mark]"
    _style_run(p.add_run(label), size=13, bold=True, color=_ACCENT)
    if text:
        qp = doc.add_paragraph()
        _add_rich_line(qp, text, size=11, color=_GREY)
        for r in qp.runs:
            r.font.italic = True


def _subheading(doc: Document, number: str, marks, text: str) -> None:
    is_or = "or" in (number or "").lower() or "অথবা" in (text or "")
    if is_or:
        or_p = doc.add_paragraph()
        or_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_run(or_p.add_run("OR"), size=12, bold=True, color=_ACCENT)

    p = doc.add_paragraph()
    label = f"({number})"
    if marks is not None:
        label += f"  [{marks} marks]" if marks != 1 else f"  [{marks} mark]"
    _style_run(p.add_run(label), size=11.5, bold=True)
    if text:
        run = p.add_run("  " + text)
        _style_run(run, size=11, italic=True, color=_GREY)


def _sources_line(doc: Document, sources: list[str]) -> None:
    if not sources:
        return
    p = doc.add_paragraph()
    _style_run(p.add_run("Sources: " + ", ".join(sources)), size=9, italic=True, color=_GREY)


def _set_columns(section, count: int, *, space: str = "360") -> None:
    """Set the Word section column count. space is twips (360 = 0.25 inch)."""
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    col_el = cols[0] if cols else OxmlElement("w:cols")
    if not cols:
        sect_pr.append(col_el)
    col_el.set(qn("w:num"), str(count))
    col_el.set(qn("w:space"), space)


def build_documents(
    paper: SolvedPaper,
    out_path: str | Path,
    *,
    include_sources: bool = True,
) -> tuple[Path, Path | None]:
    """Write the .docx, attempt a .pdf, and return (docx_path, pdf_path_or_None)."""
    out_path = Path(out_path)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(11)

    # --- header ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_run(title.add_run(paper.title or "Question Paper — Solutions"),
               size=18, bold=True, color=_ACCENT)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bits = [f"Language: {paper.language}"]
    if getattr(paper, "board", None):
        bits.append(f"Board: {paper.board}")
    if paper.class_level:
        bits.append(f"Class/Level: {paper.class_level}")
    if paper.total_marks is not None:
        bits.append(f"Total marks: {paper.total_marks}")
    _style_run(meta.add_run("   |   ".join(bits)), size=10, color=_GREY)
    doc.add_paragraph()

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    _set_columns(body_section, 2)

    # --- body ---
    for q in paper.questions:
        _heading(doc, q.number, q.marks, q.text)
        if q.subparts:
            if q.answer:
                _add_answer(doc, q.answer)
                if include_sources:
                    _sources_line(doc, q.sources)
            for sp in q.subparts:
                _subheading(doc, sp.number, sp.marks, sp.text)
                _add_answer(doc, sp.answer)
                if include_sources:
                    _sources_line(doc, sp.sources)
        else:
            _add_answer(doc, q.answer or "")
            if include_sources:
                _sources_line(doc, q.sources)
        doc.add_paragraph()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    docx_path = out_path.with_suffix(".docx")
    doc.save(str(docx_path))

    pdf_path = pdf_convert.docx_to_pdf(docx_path)
    return docx_path, pdf_path
